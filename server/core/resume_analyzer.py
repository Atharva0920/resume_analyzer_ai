"""
Enhanced AI Resume Analyzer Core Module for FastAPI Backend

This module provides comprehensive AI-powered resume analysis with:
- Multi-format file processing (PDF, DOCX, HTML, TXT)
- Industry-specific analysis and scoring
- ATS optimization scoring
- NLP-based sentiment and readability analysis
- Keyword density analysis
- Industry alignment scoring
"""

import re
import json
import logging
import asyncio
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, asdict
from datetime import datetime
import io
import tempfile
import os
from pathlib import Path

# File processing libraries
import PyPDF2
import pdfplumber
import docx
from bs4 import BeautifulSoup

# NLP and analysis libraries
import nltk
import pandas as pd
import numpy as np
from collections import Counter
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
    nltk.data.find('taggers/averaged_perceptron_tagger')
except LookupError:
    nltk.download('punkt')
    nltk.download('stopwords')
    nltk.download('averaged_perceptron_tagger')

from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.sentiment import SentimentIntensityAnalyzer

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class ContactInfo:
    """Enhanced contact information structure"""
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    address: Optional[str] = None
    linkedin: Optional[str] = None
    github: Optional[str] = None
    website: Optional[str] = None
    portfolio: Optional[str] = None


@dataclass
class Skills:
    """Categorized skills structure"""
    programming: List[str] = None
    frameworks: List[str] = None
    cloud: List[str] = None
    ai_ml: List[str] = None
    soft_skills: List[str] = None
    other_tech: List[str] = None
    
    def __post_init__(self):
        if self.programming is None:
            self.programming = []
        if self.frameworks is None:
            self.frameworks = []
        if self.cloud is None:
            self.cloud = []
        if self.ai_ml is None:
            self.ai_ml = []
        if self.soft_skills is None:
            self.soft_skills = []
        if self.other_tech is None:
            self.other_tech = []


@dataclass
class AnalysisResult:
    """Comprehensive analysis results"""
    contact_info: Dict[str, Any]
    skills: Dict[str, List[str]]
    experience_years: int
    education_level: str
    ats_score: float
    keyword_density: Dict[str, float]
    sentiment_score: float
    readability_score: float
    recommendations: List[str]
    missing_sections: List[str]
    industry_alignment: Dict[str, float]


class FileProcessor:
    """Handles file processing and text extraction"""
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.doc', '.txt', '.html', '.htm'}
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using both PyPDF2 and pdfplumber"""
        text = ""
        
        # Try pdfplumber first (better for complex layouts)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if text.strip():
                return text.strip()
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, trying PyPDF2")
        
        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise ValueError(f"Could not extract text from PDF: {e}")
        
        return text.strip()
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise ValueError(f"Could not extract text from DOCX: {e}")
    
    def extract_text_from_html(self, file_path: str) -> str:
        """Extract text from HTML file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                text = soup.get_text()
                
                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)
                
                return text
        except Exception as e:
            logger.error(f"HTML extraction failed: {e}")
            raise ValueError(f"Could not extract text from HTML: {e}")
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try different encodings
            encodings = ['latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            raise ValueError("Could not decode text file")
        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            raise ValueError(f"Could not extract text from TXT: {e}")
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Main method to extract text from any supported file format"""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        elif file_extension in ['.html', '.htm']:
            return self.extract_text_from_html(file_path)
        elif file_extension == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")


class IndustryAnalyzer:
    """Industry-specific analysis and keyword matching"""
    
    def __init__(self):
        self.industry_keywords = {
            "software_engineering": {
                "programming": [
                    "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust",
                    "scala", "kotlin", "swift", "php", "ruby", "r", "matlab", "sql"
                ],
                "frameworks": [
                    "react", "angular", "vue", "node.js", "express", "django", "flask",
                    "spring", "laravel", "rails", "asp.net", "jquery", "bootstrap"
                ],
                "cloud": [
                    "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "ansible",
                    "jenkins", "gitlab", "github actions", "circleci"
                ],
                "databases": [
                    "mysql", "postgresql", "mongodb", "redis", "elasticsearch",
                    "cassandra", "dynamodb", "sqlite", "oracle"
                ],
                "tools": [
                    "git", "jira", "confluence", "slack", "vs code", "intellij",
                    "postman", "swagger", "linux", "bash", "powershell"
                ]
            },
            "data_science": {
                "programming": [
                    "python", "r", "sql", "scala", "java", "matlab", "sas", "julia"
                ],
                "ai_ml": [
                    "machine learning", "deep learning", "tensorflow", "pytorch", "keras",
                    "scikit-learn", "pandas", "numpy", "matplotlib", "seaborn", "plotly",
                    "opencv", "nltk", "spacy", "transformers", "bert", "gpt"
                ],
                "cloud": [
                    "aws sagemaker", "azure ml", "google ai platform", "databricks",
                    "snowflake", "redshift", "bigquery"
                ],
                "tools": [
                    "jupyter", "tableau", "power bi", "looker", "apache spark",
                    "hadoop", "kafka", "airflow", "mlflow", "kubeflow"
                ]
            },
            "product_management": {
                "tools": [
                    "jira", "confluence", "asana", "trello", "notion", "figma", "sketch",
                    "invision", "miro", "lucidchart", "amplitude", "mixpanel", "google analytics"
                ],
                "methodologies": [
                    "agile", "scrum", "kanban", "lean", "design thinking", "user research",
                    "a/b testing", "mvp", "okr", "kpi", "roadmap", "backlog"
                ],
                "skills": [
                    "product strategy", "user experience", "market research", "competitive analysis",
                    "stakeholder management", "cross-functional collaboration", "data analysis"
                ]
            },
            "marketing": {
                "digital": [
                    "seo", "sem", "ppc", "google ads", "facebook ads", "linkedin ads",
                    "email marketing", "content marketing", "social media", "influencer marketing"
                ],
                "analytics": [
                    "google analytics", "facebook insights", "hootsuite", "hubspot",
                    "mailchimp", "salesforce", "marketo", "pardot"
                ],
                "skills": [
                    "brand management", "campaign management", "lead generation",
                    "conversion optimization", "marketing automation", "crm"
                ]
            },
            "finance": {
                "analysis": [
                    "financial modeling", "valuation", "dcf", "lbo", "comps",
                    "risk management", "portfolio management", "derivatives"
                ],
                "tools": [
                    "excel", "vba", "bloomberg", "reuters", "factset", "capital iq",
                    "tableau", "power bi", "python", "r", "sql"
                ],
                "certifications": [
                    "cfa", "frm", "cpa", "cma", "pmp", "six sigma"
                ]
            }
        }
    
    def get_industry_alignment_score(self, text: str, industry: str) -> Dict[str, float]:
        """Calculate alignment score with specific industry"""
        text_lower = text.lower()
        industry_data = self.industry_keywords.get(industry, {})
        
        alignment_scores = {}
        total_keywords = 0
        matched_keywords = 0
        
        for category, keywords in industry_data.items():
            category_matches = 0
            for keyword in keywords:
                if keyword.lower() in text_lower:
                    category_matches += 1
                    matched_keywords += 1
                total_keywords += 1
            
            if keywords:
                alignment_scores[category] = (category_matches / len(keywords)) * 100
        
        # Overall alignment score
        overall_score = (matched_keywords / total_keywords) * 100 if total_keywords > 0 else 0
        alignment_scores['overall'] = overall_score
        
        return alignment_scores


class ResumeParser:
    """Advanced resume parsing with NLP"""
    
    def __init__(self):
        self.stop_words = set(stopwords.words('english'))
    
    def extract_contact_info(self, text: str) -> ContactInfo:
        """Extract comprehensive contact information"""
        contact = ContactInfo()
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_matches = re.findall(email_pattern, text, re.IGNORECASE)
        if email_matches:
            contact.email = email_matches[0]
        
        # Phone patterns (multiple formats)
        phone_patterns = [
            r'(\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})',
            r'(\+?[1-9]\d{1,3}[-.\s]?)?(\d{3,4}[-.\s]?\d{3,4}[-.\s]?\d{3,4})',
        ]
        for pattern in phone_patterns:
            phone_matches = re.findall(pattern, text)
            if phone_matches:
                contact.phone = ''.join(phone_matches[0]) if isinstance(phone_matches[0], tuple) else phone_matches[0]
                break
        
        # LinkedIn
        linkedin_patterns = [
            r'linkedin\.com/in/[\w-]+',
            r'www\.linkedin\.com/in/[\w-]+',
            r'linkedin\.com/profile/view\?id=[\w-]+'
        ]
        for pattern in linkedin_patterns:
            linkedin_matches = re.findall(pattern, text, re.IGNORECASE)
            if linkedin_matches:
                contact.linkedin = linkedin_matches[0]
                break
        
        # GitHub
        github_patterns = [
            r'github\.com/[\w-]+',
            r'www\.github\.com/[\w-]+'
        ]
        for pattern in github_patterns:
            github_matches = re.findall(pattern, text, re.IGNORECASE)
            if github_matches:
                contact.github = github_matches[0]
                break
        
        # Website/Portfolio
        website_patterns = [
            r'https?://[\w.-]+\.[a-zA-Z]{2,}',
            r'www\.[\w.-]+\.[a-zA-Z]{2,}'
        ]
        for pattern in website_patterns:
            website_matches = re.findall(pattern, text, re.IGNORECASE)
            if website_matches:
                # Filter out social media links
                for url in website_matches:
                    if not any(social in url.lower() for social in ['linkedin', 'github', 'facebook', 'twitter']):
                        contact.website = url
                        break
        
        # Name extraction (improved logic)
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        for line in lines[:10]:  # Check first 10 lines
            if len(line.split()) in [2, 3, 4] and len(line) < 60:  # Reasonable name length
                # Check if line doesn't contain numbers, @ symbols, or common resume words
                if not re.search(r'[0-9@]', line) and not any(word.lower() in line.lower() 
                    for word in ['resume', 'cv', 'curriculum', 'vitae', 'address', 'email', 'phone']):
                    # Check if it looks like a name (proper case)
                    words = line.split()
                    if all(word[0].isupper() for word in words if word):
                        contact.name = line
                        break
        
        return contact
    
    def categorize_skills(self, skills_text: str, industry: str) -> Skills:
        """Categorize skills based on industry context"""
        skills = Skills()
        text_lower = skills_text.lower()
        
        # Get industry-specific categorization
        industry_analyzer = IndustryAnalyzer()
        industry_keywords = industry_analyzer.industry_keywords.get(industry, {})
        
        # Extract all potential skills
        skill_candidates = []
        
        # From comma/bullet separated lists
        skill_items = re.split(r'[,•\n\t-]', skills_text)
        for item in skill_items:
            item = item.strip()
            if item and len(item) > 1 and len(item) < 30:
                skill_candidates.append(item)
        
        # Categorize skills
        for skill in skill_candidates:
            skill_lower = skill.lower()
            
            # Check against industry categories
            for category, keywords in industry_keywords.items():
                for keyword in keywords:
                    if keyword.lower() in skill_lower or skill_lower in keyword.lower():
                        if category == 'programming':
                            skills.programming.append(skill)
                        elif category == 'frameworks':
                            skills.frameworks.append(skill)
                        elif category == 'cloud':
                            skills.cloud.append(skill)
                        elif category == 'ai_ml':
                            skills.ai_ml.append(skill)
                        elif category in ['tools', 'databases', 'digital', 'analysis']:
                            skills.other_tech.append(skill)
                        break
            else:
                # Check if it's a soft skill
                soft_skill_keywords = [
                    'leadership', 'communication', 'teamwork', 'management', 'analytical',
                    'problem solving', 'creative', 'strategic', 'collaboration'
                ]
                if any(keyword in skill_lower for keyword in soft_skill_keywords):
                    skills.soft_skills.append(skill)
                else:
                    skills.other_tech.append(skill)
        
        # Remove duplicates while preserving order
        skills.programming = list(dict.fromkeys(skills.programming))
        skills.frameworks = list(dict.fromkeys(skills.frameworks))
        skills.cloud = list(dict.fromkeys(skills.cloud))
        skills.ai_ml = list(dict.fromkeys(skills.ai_ml))
        skills.soft_skills = list(dict.fromkeys(skills.soft_skills))
        skills.other_tech = list(dict.fromkeys(skills.other_tech))
        
        return skills
    
    def extract_experience_years(self, text: str) -> int:
        """Extract years of experience from resume"""
        experience_patterns = [
            r'(\d+)\+?\s*years?\s+(?:of\s+)?experience',
            r'experience[:\s]*(\d+)\+?\s*years?',
            r'(\d+)\+?\s*years?\s+in\s+(?:the\s+)?(?:field|industry)',
            r'over\s+(\d+)\s+years?',
            r'more\s+than\s+(\d+)\s+years?'
        ]
        
        max_years = 0
        text_lower = text.lower()
        
        for pattern in experience_patterns:
            matches = re.findall(pattern, text_lower)
            for match in matches:
                try:
                    years = int(match)
                    if years <= 50:  # Reasonable upper bound
                        max_years = max(max_years, years)
                except ValueError:
                    continue
        
        # If no explicit years mentioned, try to infer from work history
        if max_years == 0:
            # Look for date ranges in experience section
            date_patterns = [
                r'(\d{4})\s*[-–]\s*(\d{4})',
                r'(\d{4})\s*[-–]\s*present',
                r'(\d{1,2}/\d{4})\s*[-–]\s*(\d{1,2}/\d{4})',
                r'(\d{1,2}/\d{4})\s*[-–]\s*present'
            ]
            
            years_found = []
            current_year = datetime.now().year
            
            for pattern in date_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                for match in matches:
                    try:
                        if isinstance(match, tuple):
                            start_year = int(match[0][:4])
                            if 'present' in match[1].lower():
                                end_year = current_year
                            else:
                                end_year = int(match[1][:4])
                            
                            if 1990 <= start_year <= current_year and start_year <= end_year:
                                years_found.append(end_year - start_year)
                    except ValueError:
                        continue
            
            if years_found:
                max_years = sum(years_found)  # Total experience
        
        return min(max_years, 50)  # Cap at 50 years
    
    def extract_education_level(self, text: str) -> str:
        """Extract highest education level"""
        text_lower = text.lower()
        
        education_levels = {
            "phd": ["ph.d", "phd", "doctorate", "doctoral"],
            "masters": ["master", "m.s", "m.a", "mba", "ms", "ma"],
            "bachelors": ["bachelor", "b.s", "b.a", "bs", "ba", "undergraduate"],
            "associates": ["associate", "a.a", "a.s", "aa", "as"],
            "high_school": ["high school", "diploma", "ged"]
        }
        
        for level, keywords in education_levels.items():
            for keyword in keywords:
                if keyword in text_lower:
                    return level
        
        return "not_specified"


class ATSScorer:
    """ATS (Applicant Tracking System) optimization scorer"""
    
    def __init__(self):
        self.ats_factors = {
            'contact_info': 15,
            'standard_sections': 20,
            'keyword_optimization': 25,
            'formatting': 15,
            'readability': 15,
            'length': 10
        }
    
    def calculate_ats_score(self, text: str, contact_info: ContactInfo, 
                           sections_found: List[str], keyword_density: Dict[str, float]) -> float:
        """Calculate comprehensive ATS score"""
        total_score = 0.0
        
        # Contact information score
        contact_score = 0
        if contact_info.email:
            contact_score += 5
        if contact_info.phone:
            contact_score += 5
        if contact_info.name:
            contact_score += 5
        total_score += min(contact_score, self.ats_factors['contact_info'])
        
        # Standard sections score
        required_sections = ['experience', 'education', 'skills']
        sections_score = (len([s for s in sections_found if any(req in s.lower() for req in required_sections)]) 
                         / len(required_sections)) * self.ats_factors['standard_sections']
        total_score += sections_score
        
        # Keyword optimization score
        avg_keyword_density = np.mean(list(keyword_density.values())) if keyword_density else 0
        keyword_score = min(avg_keyword_density * 2, self.ats_factors['keyword_optimization'])
        total_score += keyword_score
        
        # Formatting score (based on text structure)
        formatting_score = self._analyze_formatting(text)
        total_score += formatting_score
        
        # Length score
        word_count = len(text.split())
        if 300 <= word_count <= 800:  # Optimal length
            length_score = self.ats_factors['length']
        elif 200 <= word_count <= 1000:  # Acceptable length
            length_score = self.ats_factors['length'] * 0.8
        else:  # Too short or too long
            length_score = self.ats_factors['length'] * 0.5
        total_score += length_score
        
        return min(100.0, total_score)
    
    def _analyze_formatting(self, text: str) -> float:
        """Analyze text formatting for ATS compatibility"""
        score = 0.0
        
        # Check for proper section headers
        section_headers = re.findall(r'^[A-Z][A-Z\s]{3,}$', text, re.MULTILINE)
        if section_headers:
            score += 5
        
        # Check for bullet points
        bullet_points = re.findall(r'^[\s]*[•\-\*]\s+', text, re.MULTILINE)
        if bullet_points:
            score += 3
        
        # Check for consistent date formatting
        date_patterns = re.findall(r'\d{1,2}/\d{4}|\d{4}[-–]\d{4}|\w+\s+\d{4}', text)
        if date_patterns:
            score += 2
        
        # Penalize for special characters that might confuse ATS
        special_chars = len(re.findall(r'[^\w\s\-.,()/@]', text))
        if special_chars < 10:
            score += 3
        
        # Check for proper spacing
        lines = text.split('\n')
        non_empty_lines = [line for line in lines if line.strip()]
        if len(non_empty_lines) / len(lines) > 0.7:  # Good content density
            score += 2
        
        return min(score, self.ats_factors['formatting'])


class AIResumeAnalyzer:
    """Main AI-powered resume analyzer"""
    
    def __init__(self):
        self.file_processor = FileProcessor()
        self.parser = ResumeParser()
        self.industry_analyzer = IndustryAnalyzer()
        self.ats_scorer = ATSScorer()
        self.sentiment_analyzer = SentimentIntensityAnalyzer()
    
    def _calculate_readability_score(self, text: str) -> float:
        """Calculate readability score using multiple metrics"""
        sentences = sent_tokenize(text)
        words = word_tokenize(text.lower())
        
        if not sentences or not words:
            return 0.0
        
        # Basic readability metrics
        avg_sentence_length = len(words) / len(sentences)
        
        # Count syllables (approximation)
        def count_syllables(word):
            vowels = 'aeiouy'
            count = sum(1 for char in word.lower() if char in vowels)
            if word.endswith('e'):
                count -= 1
            return max(1, count)
        
        total_syllables = sum(count_syllables(word) for word in words if word.isalpha())
        avg_syllables_per_word = total_syllables / len([w for w in words if w.isalpha()])
        
        # Flesch Reading Ease approximation
        flesch_score = 206.835 - (1.015 * avg_sentence_length) - (84.6 * avg_syllables_per_word)
        
        # Convert to 0-100 scale where higher is better
        normalized_score = max(0, min(100, flesch_score))
        
        return normalized_score
    
    def _calculate_keyword_density(self, text: str, industry: str) -> Dict[str, float]:
        """Calculate keyword density for industry-specific terms"""
        text_lower = text.lower()
        word_count = len(text.split())
        
        industry_keywords = self.industry_analyzer.industry_keywords.get(industry, {})
        keyword_density = {}
        
        for category, keywords in industry_keywords.items():
            category_density = 0
            for keyword in keywords:
                count = text_lower.count(keyword.lower())
                if count > 0:
                    density = (count / word_count) * 100
                    category_density += density
            
            keyword_density[category] = round(category_density, 2)
        
        return keyword_density
    
    def _identify_missing_sections(self, text: str) -> List[str]:
        """Identify missing resume sections"""
        text_lower = text.lower()
        
        standard_sections = {
            'contact': ['contact', 'email', 'phone'],
            'summary': ['summary', 'objective', 'profile'],
            'experience': ['experience', 'work', 'employment', 'professional'],
            'education': ['education', 'degree', 'university', 'college'],
            'skills': ['skills', 'technical', 'competencies'],
            'projects': ['projects', 'portfolio'],
            'certifications': ['certification', 'certificate', 'licensed']
        }
        
        missing_sections = []
        for section, keywords in standard_sections.items():
            if not any(keyword in text_lower for keyword in keywords):
                missing_sections.append(section)
        
        return missing_sections
    
    def _generate_recommendations(self, analysis_data: Dict[str, Any]) -> List[str]:
        """Generate personalized recommendations"""
        recommendations = []
        
        # ATS Score recommendations
        if analysis_data['ats_score'] < 70:
            recommendations.append("Improve ATS compatibility by using standard section headers and bullet points")
        
        # Skills recommendations
        if len(analysis_data['skills']['programming']) < 3:
            recommendations.append("Add more technical skills relevant to your industry")
        
        # Experience recommendations
        if analysis_data['experience_years'] < 2:
            recommendations.append("Highlight internships, projects, and relevant coursework to demonstrate experience")
        
        # Contact info recommendations
        contact = analysis_data['contact_info']
        if not contact.get('linkedin'):
            recommendations.append("Add your LinkedIn profile URL to improve networking opportunities")
        
        # Missing sections recommendations
        missing = analysis_data['missing_sections']
        if 'summary' in missing:
            recommendations.append("Add a professional summary to highlight your key qualifications")
        if 'projects' in missing:
            recommendations.append("Include a projects section to showcase your practical skills")
        
        # Readability recommendations
        if analysis_data['readability_score'] < 60:
            recommendations.append("Improve readability by using shorter sentences and simpler language")
        
        # Industry alignment recommendations
        alignment = analysis_data['industry_alignment']
        if alignment.get('overall', 0) < 30:
            recommendations.append("Include more industry-specific keywords and technologies")
        
        # Keyword density recommendations
        keyword_density = analysis_data['keyword_density']
        low_categories = [cat for cat, density in keyword_density.items() if density < 1.0]
        if low_categories:
            recommendations.append(f"Strengthen these skill areas: {', '.join(low_categories[:3])}")
        
        return recommendations[:8]  # Limit to most important recommendations
    
    async def analyze_resume(self, text: str, industry: str = "software_engineering") -> AnalysisResult:
        """Main analysis method - now async for better performance"""
        try:
            # Extract basic information
            contact_info = self.parser.extract_contact_info(text)
            
            # Extract skills section text for detailed analysis
            skills_section_match = re.search(
                r'(?:skills?|technical skills?|competencies)[:\s]*(.*?)(?=\n\n|\n[A-Z]|$)', 
                text, re.IGNORECASE | re.DOTALL
            )
            skills_text = skills_section_match.group(1) if skills_section_match else text
            
            # Categorize skills
            categorized_skills = self.parser.categorize_skills(skills_text, industry)
            
            # Extract experience and education
            experience_years = self.parser.extract_experience_years(text)
            education_level = self.parser.extract_education_level(text)
            
            # Calculate scores
            ats_score = self.ats_scorer.calculate_ats_score(
                text, contact_info, 
                ['experience', 'education', 'skills'],  # Simplified for now
                self._calculate_keyword_density(text, industry)
            )
            
            # NLP analysis
            sentiment_scores = self.sentiment_analyzer.polarity_scores(text)
            sentiment_score = max(0, sentiment_scores['compound'] + 1) / 2  # Normalize to 0-1
            
            readability_score = self._calculate_readability_score(text)
            keyword_density = self._calculate_keyword_density(text, industry)
            industry_alignment = self.industry_analyzer.get_industry_alignment_score(text, industry)
            missing_sections = self._identify_missing_sections(text)
            
            # Prepare analysis data for recommendations
            analysis_data = {
                'contact_info': asdict(contact_info),
                'skills': asdict(categorized_skills),
                'experience_years': experience_years,
                'education_level': education_level,
                'ats_score': ats_score,
                'keyword_density': keyword_density,
                'sentiment_score': sentiment_score,
                'readability_score': readability_score,
                'missing_sections': missing_sections,
                'industry_alignment': industry_alignment
            }
            
            recommendations = self._generate_recommendations(analysis_data)
            
            return AnalysisResult(
                contact_info=asdict(contact_info),
                skills=asdict(categorized_skills),
                experience_years=experience_years,
                education_level=education_level,
                ats_score=round(ats_score, 1),
                keyword_density=keyword_density,
                sentiment_score=round(sentiment_score, 3),
                readability_score=round(readability_score, 1),
                recommendations=recommendations,
                missing_sections=missing_sections,
                industry_alignment=industry_alignment
            )
            
        except Exception as e:
            logger.error(f"Analysis failed: {str(e)}")
            raise Exception(f"Resume analysis failed: {str(e)}")
    
    def analyze_resume_sync(self, text: str, industry: str = "software_engineering") -> AnalysisResult:
        """Synchronous version of analyze_resume for backwards compatibility"""
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            return loop.run_until_complete(self.analyze_resume(text, industry))
        finally:
            loop.close()


# Additional utility functions for FastAPI integration

async def analyze_resume_from_text(text: str, industry: str = "software_engineering") -> Dict[str, Any]:
    """Convenience function for text analysis"""
    analyzer = AIResumeAnalyzer()
    result = await analyzer.analyze_resume(text, industry)
    return asdict(result)


async def analyze_resume_from_file(file_path: str, industry: str = "software_engineering") -> Dict[str, Any]:
    """Convenience function for file analysis"""
    analyzer = AIResumeAnalyzer()
    
    # Extract text from file
    text = analyzer.file_processor.extract_text_from_file(file_path)
    
    # Analyze the text
    result = await analyzer.analyze_resume(text, industry)
    
    return {
        **asdict(result),
        'extracted_text_length': len(text),
        'file_processed': True
    }


class ResumeComparer:
    """Compare multiple resumes or resume versions"""
    
    def __init__(self):
        self.analyzer = AIResumeAnalyzer()
    
    async def compare_resumes(self, resume_texts: List[str], industry: str = "software_engineering") -> Dict[str, Any]:
        """Compare multiple resumes"""
        if len(resume_texts) < 2:
            raise ValueError("At least 2 resumes required for comparison")
        
        analyses = []
        for i, text in enumerate(resume_texts):
            result = await self.analyzer.analyze_resume(text, industry)
            analyses.append({
                'resume_id': i + 1,
                'analysis': asdict(result)
            })
        
        # Calculate comparison metrics
        comparison = {
            'total_resumes': len(resume_texts),
            'analyses': analyses,
            'comparison_metrics': {
                'highest_ats_score': max(a['analysis']['ats_score'] for a in analyses),
                'average_experience_years': np.mean([a['analysis']['experience_years'] for a in analyses]),
                'best_industry_alignment': max(a['analysis']['industry_alignment']['overall'] for a in analyses),
                'most_comprehensive_skills': max(len(a['analysis']['skills']['programming']) + 
                                               len(a['analysis']['skills']['frameworks']) for a in analyses)
            },
            'recommendations': {
                'best_overall_resume': max(analyses, key=lambda x: x['analysis']['ats_score'])['resume_id'],
                'areas_for_improvement': self._get_comparison_recommendations(analyses)
            }
        }
        
        return comparison
    
    def _get_comparison_recommendations(self, analyses: List[Dict]) -> List[str]:
        """Generate recommendations based on resume comparison"""
        recommendations = []
        
        # Find best practices from top-performing resume
        best_resume = max(analyses, key=lambda x: x['analysis']['ats_score'])
        best_analysis = best_resume['analysis']
        
        recommendations.append(f"Resume {best_resume['resume_id']} has the highest ATS score ({best_analysis['ats_score']})")
        
        # Identify common weaknesses
        common_missing = {}
        for analysis in analyses:
            for missing in analysis['analysis']['missing_sections']:
                common_missing[missing] = common_missing.get(missing, 0) + 1
        
        if common_missing:
            most_common_missing = max(common_missing.items(), key=lambda x: x[1])
            if most_common_missing[1] > len(analyses) / 2:
                recommendations.append(f"Most resumes are missing: {most_common_missing[0]}")
        
        return recommendations


class ResumeOptimizer:
    """Optimize resume for specific job descriptions"""
    
    def __init__(self):
        self.analyzer = AIResumeAnalyzer()
        self.vectorizer = TfidfVectorizer(stop_words='english', max_features=1000)
    
    async def optimize_for_job(self, resume_text: str, job_description: str, 
                              industry: str = "software_engineering") -> Dict[str, Any]:
        """Optimize resume for specific job description"""
        try:
            # Analyze current resume
            current_analysis = await self.analyzer.analyze_resume(resume_text, industry)
            
            # Extract keywords from job description
            job_keywords = self._extract_job_keywords(job_description, industry)
            
            # Calculate similarity score
            similarity_score = self._calculate_similarity(resume_text, job_description)
            
            # Generate optimization recommendations
            optimization_recommendations = self._generate_optimization_recommendations(
                current_analysis, job_keywords, job_description
            )
            
            return {
                'current_analysis': asdict(current_analysis),
                'job_match_score': round(similarity_score * 100, 1),
                'missing_keywords': job_keywords['missing'],
                'matching_keywords': job_keywords['matching'],
                'optimization_recommendations': optimization_recommendations,
                'priority_improvements': optimization_recommendations[:5]
            }
            
        except Exception as e:
            logger.error(f"Optimization failed: {str(e)}")
            raise Exception(f"Resume optimization failed: {str(e)}")
    
    def _extract_job_keywords(self, job_description: str, industry: str) -> Dict[str, List[str]]:
        """Extract relevant keywords from job description"""
        job_text_lower = job_description.lower()
        
        # Get industry-specific keywords
        industry_analyzer = IndustryAnalyzer()
        industry_keywords = industry_analyzer.industry_keywords.get(industry, {})
        
        matching_keywords = []
        all_industry_keywords = []
        
        for category, keywords in industry_keywords.items():
            all_industry_keywords.extend(keywords)
            for keyword in keywords:
                if keyword.lower() in job_text_lower:
                    matching_keywords.append(keyword)
        
        # Find missing keywords that are common in the industry
        missing_keywords = [kw for kw in all_industry_keywords if kw not in matching_keywords]
        
        return {
            'matching': matching_keywords,
            'missing': missing_keywords[:20],  # Top 20 missing keywords
            'total_found': len(matching_keywords)
        }
    
    def _calculate_similarity(self, resume_text: str, job_description: str) -> float:
        """Calculate similarity between resume and job description"""
        try:
            # Fit vectorizer on both documents
            documents = [resume_text, job_description]
            tfidf_matrix = self.vectorizer.fit_transform(documents)
            
            # Calculate cosine similarity
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            
            return similarity
        except Exception:
            return 0.0
    
    def _generate_optimization_recommendations(self, current_analysis: AnalysisResult, 
                                             job_keywords: Dict[str, List[str]], 
                                             job_description: str) -> List[str]:
        """Generate specific optimization recommendations"""
        recommendations = []
        
        # Keyword optimization
        if job_keywords['missing']:
            recommendations.append(f"Add these job-relevant keywords: {', '.join(job_keywords['missing'][:5])}")
        
        # Skills gap analysis
        job_skills = set(kw.lower() for kw in job_keywords['matching'])
        resume_skills = set()
        
        # Collect all resume skills
        for skill_category in current_analysis.skills.values():
            if isinstance(skill_category, list):
                resume_skills.update(skill.lower() for skill in skill_category)
        
        skill_gaps = job_skills - resume_skills
        if skill_gaps:
            recommendations.append(f"Highlight these skills if you have them: {', '.join(list(skill_gaps)[:3])}")
        
        # Experience optimization
        if current_analysis.experience_years < 3:
            recommendations.append("Emphasize projects, internships, and relevant coursework to demonstrate experience")
        
        # ATS optimization
        if current_analysis.ats_score < 80:
            recommendations.append("Improve ATS compatibility with better formatting and standard section headers")
        
        # Industry alignment
        if current_analysis.industry_alignment.get('overall', 0) < 50:
            recommendations.append("Better align your resume with industry-specific terminology and requirements")
        
        return recommendations


# Example usage and testing functions
if __name__ == "__main__":
    import asyncio
    
    async def test_analyzer():
        """Test the analyzer with sample data"""
        analyzer = AIResumeAnalyzer()
        
        sample_resume = """
        John Doe
        john.doe@email.com | (555) 123-4567 | linkedin.com/in/johndoe | github.com/johndoe
        
        PROFESSIONAL SUMMARY
        Experienced Software Engineer with 5 years of experience in Python, JavaScript, and cloud technologies.
        Passionate about building scalable web applications and leading development teams.
        
        TECHNICAL SKILLS
        Programming Languages: Python, JavaScript, TypeScript, Java, SQL
        Web Technologies: React, Node.js, Express, Django, Flask
        Cloud & DevOps: AWS, Docker, Kubernetes, Jenkins, Terraform
        Databases: PostgreSQL, MongoDB, Redis
        Tools: Git, JIRA, VS Code, Postman
        
        PROFESSIONAL EXPERIENCE
        Senior Software Engineer | TechCorp Inc. | Jan 2020 - Present
        • Led development of microservices architecture serving 1M+ users
        • Implemented CI/CD pipelines reducing deployment time by 60%
        • Mentored 3 junior developers and conducted code reviews
        • Built RESTful APIs using Python and Django
        
        Software Engineer | StartupXYZ | Jun 2018 - Dec 2019
        • Developed responsive web applications using React and Node.js
        • Optimized database queries improving performance by 40%
        • Collaborated with cross-functional teams using Agile methodologies
        
        EDUCATION
        Bachelor of Science in Computer Science
        University of Technology | 2018
        
        PROJECTS
        E-commerce Platform | Python, Django, PostgreSQL, AWS
        • Built full-stack e-commerce application with payment integration
        • Implemented user authentication and order management system
        
        CERTIFICATIONS
        AWS Certified Solutions Architect
        """
        
        # Test basic analysis
        print("Testing basic resume analysis...")
        result = await analyzer.analyze_resume(sample_resume, "software_engineering")
        
        print(f"ATS Score: {result.ats_score}")
        print(f"Experience Years: {result.experience_years}")
        print(f"Education Level: {result.education_level}")
        print(f"Contact Email: {result.contact_info.get('email', 'Not found')}")
        print(f"Programming Skills: {result.skills.get('programming', [])}")
        print(f"Recommendations: {len(result.recommendations)}")
        print("\nTop Recommendations:")
        for i, rec in enumerate(result.recommendations[:3], 1):
            print(f"{i}. {rec}")
        
        # Test job optimization
        print("\n" + "="*50)
        print("Testing job optimization...")
        
        job_description = """
        We are seeking a Senior Full Stack Developer with expertise in Python, React, and AWS.
        The ideal candidate should have experience with microservices, containerization with Docker,
        and cloud deployment. Knowledge of machine learning and data science is a plus.
        Strong leadership skills and experience mentoring junior developers required.
        """
        
        optimizer = ResumeOptimizer()
        optimization_result = await optimizer.optimize_for_job(
            sample_resume, job_description, "software_engineering"
        )
        
        print(f"Job Match Score: {optimization_result['job_match_score']}%")
        print(f"Missing Keywords: {optimization_result['missing_keywords'][:5]}")
        print(f"Matching Keywords: {len(optimization_result['matching_keywords'])}")
        print("\nOptimization Recommendations:")
        for i, rec in enumerate(optimization_result['priority_improvements'], 1):
            print(f"{i}. {rec}")
    
    # Run the test
    asyncio.run(test_analyzer())